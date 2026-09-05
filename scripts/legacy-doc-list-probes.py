#!/usr/bin/env python3
"""Generate passive, parameter-controlled Office experiments, not golden files.

Requires python-docx. Run against a NEW local output directory. No Office app
automation or converter policy is embedded here. The saved binary must be
reopened in Word before exporting its PDF and must be inspected to establish
which source parameters survived conversion. See MS-DOC 2.4.6.3/2.6.2 and
ECMA-376 17.3.1.12, 17.3.1.38, 17.7.2 and 17.9.
"""
import argparse
from copy import deepcopy
from datetime import datetime, timezone
from hashlib import sha256
from io import BytesIO
import json
import re
from pathlib import Path
from zipfile import ZipFile, ZipInfo, ZIP_DEFLATED

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def matrix(phase="baseline"):
    if phase == "interactions":
        return interaction_matrix()
    if phase == "style-association":
        return style_association_matrix()
    if phase == "bidi-boundaries":
        return bidi_boundary_matrix()
    if phase != "baseline":
        raise ValueError("unknown experiment phase")
    cases = []

    def add(parent, **changes):
        values = deepcopy(parent["parameters"])
        values.update(changes)
        case = {"id": f"P{len(cases) + 1:03}", "parent": parent["id"],
                "changed": list(changes), "parameters": values}
        cases.append(case)
        return case

    for rtl in [False, True]:
        base = {"id": f"P{len(cases) + 1:03}", "parent": None, "changed": [],
                "parameters": dict(rtl=rtl, list_left=720, list_right=0,
                    list_first=-360, direct_left=None, direct_right=None,
                    direct_first=None, empty_direct_ind=False, style_left=None,
                    style_first=None, numbering_in_style=False, linked_style=False,
                    suffix="space", list_tab=None, direct_tab=None, list_bidi=None)}
        cases.append(base)
        for key, values in [
            ("direct_left", [-360, 0, 720, 1440]),
            ("direct_right", [-360, 0, 720]),
            ("direct_first", [-720, -360, 0, 360, 720]),
            ("list_left", [0, 1440]),
            ("list_first", [0, 360]),
            ("list_bidi", [False, True]),
        ]:
            for value in values:
                add(base, **{key: value})
        add(base, empty_direct_ind=True)
        styled = add(base, style_left=1440)
        add(styled, direct_left=0)
        add(styled, linked_style=True)
        styled_first = add(base, style_first=720)
        add(styled_first, direct_first=0)
        inherited = add(base, numbering_in_style=True)
        add(inherited, direct_left=0)
        tabs = add(base, suffix="tab")
        tab_stop = add(tabs, list_tab=720)
        add(tab_stop, direct_tab=1440)
        add(tab_stop, direct_tab="clear")
        # An unchanged repeat detects order-dependent conversion/export drift.
        add(base)
    return cases


def interaction_matrix():
    """Conflict and one-twip boundaries; outcomes must be measured in Word."""
    cases = []

    def add(parent, **changes):
        values = {**parent["parameters"], **changes}
        case = {"id": f"Q{len(cases) + 1:03}", "parent": parent["id"],
                "changed": list(changes), "parameters": values}
        cases.append(case)
        return case

    for original in (case for case in matrix() if case["parent"] is None):
        base = {"id": f"Q{len(cases) + 1:03}", "parent": None, "changed": [],
                "parameters": {**original["parameters"], "list_right": 720}}
        cases.append(base)
        for key, values in [
            ("direct_right", [-720, -1, 0, 1, 719, 720, 721, 1440]),
            ("direct_left", [-1, 0, 1, 719, 720, 721]),
            ("direct_first", [-720, -361, -360, -359, -1, 0, 1, 359, 360, 361]),
        ]:
            for value in values:
                add(base, **{key: value})
        left = add(base, direct_left=0)
        both = add(left, direct_right=0)
        add(both, direct_first=0)
        bidi = add(base, list_bidi=not base["parameters"]["rtl"])
        add(bidi, direct_left=0)
        add(bidi, direct_right=0)
        add(base)
    return cases


def element(name, **attributes):
    result = OxmlElement("w:" + name)
    for key, value in attributes.items():
        result.set(qn("w:" + key), str(value))
    return result


def style_association_matrix():
    """Try to separate paragraph-style association from direct indentation.

    Word may still associate all saved lists with a style. Only inspection of
    the saved LSTF.rgistdPara establishes whether an unlinked control exists.
    """
    cases = []

    def add(parent, **changes):
        case = {"id": f"R{len(cases) + 1:03}", "parent": parent["id"],
                "changed": list(changes),
                "parameters": {**parent["parameters"], **changes}}
        cases.append(case)
        return case

    for original in (case for case in matrix() if case["parent"] is None):
        base = {"id": f"R{len(cases) + 1:03}", "parent": None, "changed": [],
                "parameters": {**original["parameters"], "paragraph_styles": "uniform"}}
        cases.append(base)
        groups = [base] + [add(base, paragraph_styles=mode)
                           for mode in ["normal", "mixed-normal", "alternating"]]
        for group in groups:
            left = add(group, direct_left=0)
            add(group, direct_left=1440)
            add(group, direct_first=0)
            add(left, direct_first=0)
            add(group)
    return cases


def add_indent(parent, left=None, right=None, first=None, empty=False):
    if not empty and left is None and right is None and first is None:
        return
    attrs = {}
    if left is not None:
        attrs["left"] = left
    if right is not None:
        attrs["right"] = right
    if first is not None:
        attrs["hanging" if first < 0 else "firstLine"] = abs(first)
    parent.append(element("ind", **attrs))


def bidi_boundary_matrix():
    """Separate punctuation/run boundaries from list indentation.

    Strong Latin text with w:rtl=true is explicitly unspecified by ECMA-376
    17.3.2.30, so these controls use only absent or explicitly false run rtl.
    """
    cases = []

    def add(parent, **changes):
        case = {"id": f"S{len(cases) + 1:03}", "parent": parent["id"],
                "changed": list(changes), "parameters": {**parent["parameters"], **changes}}
        cases.append(case)
        return case

    for original in (case for case in matrix() if case["parent"] is None):
        base = {"id": f"S{len(cases) + 1:03}", "parent": None, "changed": [],
                "parameters": {**original["parameters"], "terminal_punctuation": ".",
                               "text_runs": "whole", "run_rtl": None, "numbered": True}}
        cases.append(base)
        punctuation = {value: add(base, terminal_punctuation=value) for value in ["", "!", ":", "?"]}
        for group in [base, punctuation["!"]]:
            for mode in ["punctuation", "words"]:
                add(group, text_runs=mode)
        plain = add(base, numbered=False)
        add(plain, terminal_punctuation="")
        add(plain, text_runs="punctuation")
        explicit_off = add(base, run_rtl=False)
        add(explicit_off, text_runs="punctuation")
        add(base)
        add(plain)
    return cases


def add_bidi_probe_text(paragraph, label, params):
    terminal = params["terminal_punctuation"]
    mode = params["text_runs"]
    if terminal not in ["", ".", "!", ":", "?"] or mode not in ["whole", "punctuation", "words"]:
        raise ValueError("unknown bidi boundary condition")
    if params["run_rtl"] not in [None, False]:
        raise ValueError("strong Latin rtl=true is not a defined-behavior control")
    texts = [label + " marker line alpha bravo charlie delta",
             "Continuation line alpha bravo charlie delta echo foxtrot golf hotel india juliet kilo lima mike november oscar papa"]
    for index, text in enumerate(texts):
        if index:
            paragraph.add_run().add_break()
        if mode == "punctuation" and terminal:
            pieces = [text, terminal]
        elif mode == "words":
            pieces = re.findall(r"\S+\s*|\s+", text + terminal)
        else:
            pieces = [text + terminal]
        for piece in pieces:
            run = paragraph.add_run(piece)
            if params["run_rtl"] is False:
                run._r.get_or_add_rPr().append(element("rtl", val=0))


def add_tabs(parent, value):
    if value is not None:
        tabs = element("tabs")
        tabs.append(element("tab", val="clear" if value == "clear" else "num",
                            pos=720 if value == "clear" else value))
        parent.append(tabs)


def add_reference(parent, identity):
    ref = element("numPr")
    ref.append(element("ilvl", val=0))
    ref.append(element("numId", val=identity))
    parent.append(ref)


def build(cases):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Arial", Pt(12)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = Pt(15)
    title = doc.styles["Title"]
    title.font.name, title.font.size = "Arial", Pt(18)
    title.font.color.rgb = normal.font.color.rgb = RGBColor(0, 0, 0)
    boundary_experiment = any("terminal_punctuation" in case["parameters"] for case in cases)
    if boundary_experiment:
        # Keep P/Q/R bytes unchanged; the new probe title has no decorative rule.
        for border in title.element.xpath("./w:pPr/w:pBdr"):
            border.getparent().remove(border)
    for style in [normal, title]:
        fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        for field in ["ascii", "hAnsi", "eastAsia", "cs"]:
            fonts.set(qn("w:" + field), "Arial")
    numbering = doc.part.numbering_part.element
    for child in list(numbering):
        numbering.remove(child)
    definitions, instances = [], []
    for index, case in enumerate(cases):
        p = case["parameters"]
        identity = index + 1
        style_mode = p.get("paragraph_styles", "uniform")
        if style_mode not in ["uniform", "normal", "mixed-normal", "alternating"]:
            raise ValueError("unknown paragraph style pattern")
        style = normal if style_mode == "normal" else doc.styles.add_style(
            "Probe" + case["id"], WD_STYLE_TYPE.PARAGRAPH)
        if style is not normal:
            style.base_style = normal
        style_ppr = style.element.get_or_add_pPr()
        if p["numbering_in_style"]:
            add_reference(style_ppr, identity)
        add_indent(style_ppr, left=p["style_left"], first=p["style_first"])
        alternate = None
        if style_mode == "alternating":
            alternate = doc.styles.add_style("Probe" + case["id"] + "Alternate", WD_STYLE_TYPE.PARAGRAPH)
            alternate.base_style = normal
            alternate.element.append(deepcopy(style_ppr))
        abstract = element("abstractNum", abstractNumId=identity)
        abstract.append(element("multiLevelType", val="singleLevel"))
        level = element("lvl", ilvl=0)
        level.append(element("start", val=1))
        level.append(element("numFmt", val="decimal"))
        if p["linked_style"]:
            level.append(element("pStyle", val=style.style_id))
        level.append(element("suff", val=p["suffix"]))
        level.append(element("lvlText", val="%1."))
        level.append(element("lvlJc", val="left"))
        level_ppr = element("pPr")
        add_tabs(level_ppr, p["list_tab"])
        if p["list_bidi"] is not None:
            level_ppr.append(element("bidi", val=int(p["list_bidi"])))
        add_indent(level_ppr, p["list_left"], p["list_right"], p["list_first"])
        level.append(level_ppr)
        abstract.append(level)
        definitions.append(abstract)
        instance = element("num", numId=identity)
        instance.append(element("abstractNumId", val=identity))
        instances.append(instance)
        heading = doc.add_paragraph(("Bidirectional text probe " if boundary_experiment else "List indentation probe ") + case["id"], "Title")
        heading.paragraph_format.page_break_before = index > 0
        doc.add_paragraph("Compare word order and punctuation before and after saving as DOC."
                          if boundary_experiment else "Compare the marker, first line, continuation line and wrapped text.")
        for label in ["First", "Second"]:
            paragraph_style = style
            if label == "Second" and style_mode == "mixed-normal":
                paragraph_style = normal
            elif label == "Second" and style_mode == "alternating":
                paragraph_style = alternate
            paragraph = doc.add_paragraph(style=paragraph_style)
            props = paragraph._p.get_or_add_pPr()
            if not p["numbering_in_style"] and p.get("numbered", True):
                add_reference(props, identity)
            add_tabs(props, p["direct_tab"])
            props.append(element("bidi", val=int(p["rtl"])))
            add_indent(props, p["direct_left"], p["direct_right"], p["direct_first"], p["empty_direct_ind"])
            if "terminal_punctuation" in p:
                add_bidi_probe_text(paragraph, label, p)
            else:
                run = paragraph.add_run(label + " marker line alpha bravo charlie delta.")
                run.add_break()
                run.add_text("Continuation line alpha bravo charlie delta echo foxtrot golf hotel india juliet kilo lima mike november oscar papa.")
    numbering.extend(definitions + instances)
    doc.core_properties.author = doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "Bidirectional text probes" if boundary_experiment else "List indentation probes"
    doc.core_properties.created = doc.core_properties.modified = datetime(2000, 1, 1, tzinfo=timezone.utc)
    raw = BytesIO()
    doc.save(raw)
    output = BytesIO()
    # Stable archive timestamps make reruns and input identity auditable.
    with ZipFile(raw) as source, ZipFile(output, "w", ZIP_DEFLATED) as destination:
        for name in sorted(source.namelist()):
            info = ZipInfo(name, (2000, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            destination.writestr(info, source.read(name))
    return output.getvalue()


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", type=Path, help="new directory; existing paths are rejected")
    parser.add_argument("--phase", choices=["baseline", "interactions", "style-association", "bidi-boundaries"], default="baseline")
    args = parser.parse_args()
    cases = matrix(args.phase)
    payload = build(cases)
    args.output.mkdir(parents=True, exist_ok=False)
    (args.output / "list-indent-probes.docx").write_bytes(payload)
    manifest = {"phase": "authored-not-office-verified", "experiment": args.phase,
                "inputSha256": sha256(payload).hexdigest(),
                "units": "twips", "cases": cases,
                "oracle": "Word PDF from a saved and reopened DOC; not the source DOCX PDF",
                "inference": "none; inspect saved binary properties before classifying outcomes"}
    (args.output / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n")
    print(f"Created {len(cases)} controlled cases; Office conversion and rendering remain required")


if __name__ == "__main__":
    main()

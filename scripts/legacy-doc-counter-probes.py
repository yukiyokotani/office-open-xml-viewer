#!/usr/bin/env python3
"""Generate passive DOCX controls for numbering-instance counter identity.

The source DOCX and its local source PDF are the primary oracle for the emitted
OOXML behavior. DOC down-save, direct-DOC PDF, and roundtrip DOCX are separate
observations of the binary conversion. No counter outcome is encoded here.
"""
import argparse
from datetime import datetime, timezone
from hashlib import sha256
from io import BytesIO
import json
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def element(name, **attributes):
    node = OxmlElement("w:" + name)
    for key, value in attributes.items():
        node.set(qn("w:" + key), str(value))
    return node


def cases():
    definitions = [
        ("same instance baseline", "1111", {"alias": "none"}),
        ("bare aliases alternating", "1212", {"alias": "bare"}),
        ("bare aliases in blocks", "1122", {"alias": "bare"}),
        ("identical full level replacement", "1212", {"alias": "full-level-identical"}),
        ("bold marker full level replacement", "1212", {"alias": "full-level-marker-bold"}),
        ("start override seven", "1212", {"alias": "start-override", "startOverride": 7}),
        ("replacement start seven", "1212", {"alias": "full-level-start", "levelStart": 7}),
        ("independent equivalent abstracts", "1212", {"alias": "independent-abstract"}),
        ("repeat bare aliases alternating", "1212", {"alias": "bare", "repeatOf": "C02"}),
    ]
    return [{"id": f"C{index:02}", "title": title, "numPattern": pattern,
             "settings": settings, "counterOutcome": "UNOBSERVED"}
            for index, (title, pattern, settings) in enumerate(definitions, 1)]


def level(start=1, bold=False):
    node = element("lvl", ilvl=0)
    node.append(element("start", val=start))
    node.append(element("numFmt", val="decimal"))
    node.append(element("suff", val="space"))
    node.append(element("lvlText", val="%1."))
    node.append(element("lvlJc", val="left"))
    ppr = element("pPr")
    tabs = element("tabs")
    tabs.append(element("tab", val="num", pos=720))
    ppr.append(tabs)
    ppr.append(element("ind", left=720, hanging=360))
    node.append(ppr)
    if bold:
        rpr = element("rPr")
        rpr.append(element("b"))
        node.append(rpr)
    return node


def abstract(identity):
    node = element("abstractNum", abstractNumId=identity)
    node.append(element("multiLevelType", val="singleLevel"))
    node.append(level())
    return node


def instance(identity, abstract_id, settings, second):
    node = element("num", numId=identity)
    node.append(element("abstractNumId", val=abstract_id))
    if second:
        mode = settings["alias"]
        if mode == "full-level-identical":
            override = element("lvlOverride", ilvl=0)
            override.append(level())
            node.append(override)
        elif mode == "full-level-marker-bold":
            override = element("lvlOverride", ilvl=0)
            override.append(level(bold=True))
            node.append(override)
        elif mode == "start-override":
            override = element("lvlOverride", ilvl=0)
            override.append(element("startOverride", val=settings["startOverride"]))
            node.append(override)
        elif mode == "full-level-start":
            override = element("lvlOverride", ilvl=0)
            override.append(level(start=settings["levelStart"]))
            node.append(override)
    return node


def reference(paragraph, identity):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = element("numPr")
    num_pr.append(element("ilvl", val=0))
    num_pr.append(element("numId", val=identity))
    ppr.append(num_pr)


def validate_case(case):
    required = {"id", "title", "numPattern", "settings", "counterOutcome"}
    if set(case) != required or case["counterOutcome"] != "UNOBSERVED":
        raise ValueError("invalid counter probe case")
    if len(case["numPattern"]) != 4 or set(case["numPattern"]) - {"1", "2"}:
        raise ValueError("invalid counter probe instance pattern")
    settings = case["settings"]
    mode = settings.get("alias")
    keys = {
        "none": {"alias"},
        "bare": {"alias"} | ({"repeatOf"} if "repeatOf" in settings else set()),
        "full-level-identical": {"alias"},
        "full-level-marker-bold": {"alias"},
        "start-override": {"alias", "startOverride"},
        "full-level-start": {"alias", "levelStart"},
        "independent-abstract": {"alias"},
    }.get(mode)
    if keys is None or set(settings) != keys:
        raise ValueError("unsupported counter probe settings")
    if mode == "none" and case["numPattern"] != "1111":
        raise ValueError("baseline cannot reference its unused alias")
    for key in ("startOverride", "levelStart"):
        if key in settings and (not isinstance(settings[key], int) or settings[key] < 0):
            raise ValueError("invalid counter probe start")


def build(matrix):
    if len({case.get("id") for case in matrix}) != len(matrix):
        raise ValueError("duplicate counter probe case ID")
    for case in matrix:
        validate_case(case)
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    normal, title = doc.styles["Normal"], doc.styles["Title"]
    normal.font.name, normal.font.size = "Arial", Pt(12)
    title.font.name, title.font.size = "Arial", Pt(18)
    normal.font.color.rgb = title.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = normal.paragraph_format.space_after = Pt(0)
    for border in title.element.xpath("./w:pPr/w:pBdr"):
        border.getparent().remove(border)
    for style in (normal, title):
        fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        for field in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn("w:" + field), "Arial")

    numbering = doc.part.numbering_part.element
    for child in list(numbering):
        numbering.remove(child)
    abstracts, instances = [], []
    for index, case in enumerate(matrix, 1):
        abstract_id = index * 10
        second_abstract = abstract_id if case["settings"]["alias"] != "independent-abstract" else abstract_id + 1
        abstracts.append(abstract(abstract_id))
        if second_abstract != abstract_id:
            abstracts.append(abstract(second_abstract))
        first_num, second_num = index * 2 - 1, index * 2
        instances.append(instance(first_num, abstract_id, case["settings"], False))
        instances.append(instance(second_num, second_abstract, case["settings"], True))

        heading = doc.add_paragraph(f"Counter instance probe {case['id']}", "Title")
        heading.paragraph_format.page_break_before = index > 1
        doc.add_paragraph(case["title"] + f". Instance sequence {case['numPattern']}. Counter outcome is unobserved until native Office review.")
        for offset, selector in enumerate(case["numPattern"]):
            paragraph = doc.add_paragraph()
            reference(paragraph, first_num if selector == "1" else second_num)
            paragraph.add_run(f"Item {'ABCD'[offset]}")
    numbering.extend(abstracts + instances)

    doc.core_properties.author = doc.core_properties.last_modified_by = ""
    doc.core_properties.title = "Counter instance probes"
    doc.core_properties.created = doc.core_properties.modified = datetime(2000, 1, 1, tzinfo=timezone.utc)
    raw, output = BytesIO(), BytesIO()
    doc.save(raw)
    with ZipFile(raw) as source, ZipFile(output, "w", ZIP_DEFLATED) as destination:
        for name in sorted(source.namelist()):
            info = ZipInfo(name, (2000, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            destination.writestr(info, source.read(name))
    return output.getvalue()


def manifest(payload, matrix):
    return {
        "phase": "authored-not-office-verified",
        "inputSha256": sha256(payload).hexdigest(),
        "oracle": "source-docx-pdf",
        "oraclePolicy": "Source DOCX PDF establishes emitted OOXML behavior; DOC down-save, direct-DOC PDF, and roundtrip DOCX are separate binary observations.",
        "cases": matrix,
    }


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", type=Path, help="new directory; existing paths are rejected")
    args = parser.parse_args()
    matrix = cases()
    payload = build(matrix)
    args.output.mkdir(parents=True, exist_ok=False)
    (args.output / "counter-instance-probes.docx").write_bytes(payload)
    (args.output / "manifest.json").write_text(json.dumps(manifest(payload, matrix), indent=2) + "\n")
    print("Created 9 counter controls; every counter outcome remains UNOBSERVED")


if __name__ == "__main__":
    main()
